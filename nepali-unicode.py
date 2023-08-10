from docx import Document
array_one = [
    "‘", "?", "क़", "ख़", "ग़", "ज़", "ड़", "ढ़", "फ़", "ॐ", "ऽ", "।", "m'", "m]",
    "mfF", "mF", "०", "१", "२", "३", "४", "५", "६", "७", "८", "९", "फ्र", "झ", "फ",
    "क्त", "क्र", "ल", "ज्ञ्", "द्घ", "ज्ञ", "द्द", "द्ध", "श्र", "रु", "द्य", "क्ष्",
    "क्ष", "त्त", "द्म", "त्र", "ध्र", "ङ्घ", "ड्ड", "द्र", "ट्ट", "ड्ढ", "ठ्ठ", "रू",
    "हृ", "ङ्ग", "त्र", "ङ्क", "ङ्ख", "ट्ठ", "द्व", "ट्र", "ठ्र", "ड्र", "ढ्र", "्र",
    "ड़", "ढ़", "क्", "क", "ख्", "ख", "ग्", "ग", "घ्", "घ", "ङ", "च्", "च", "छ", "ज्",
    "ज", "झ्", "झ", "ञ्", "ञ", "ट", "ठ", "ड", "ढ", "ण्", "ण", "त्", "त", "थ्", "थ", "द",
    "ध्", "ध", "न्", "न", "प्", "प", "फ्", "ब्", "ब", "भ्", "भ", "म्", "म", "य", "र", "ल्",
    "ल", "व्", "व", "श्", "श", "ष्", "ष", "स्", "स", "ह्", "ह", "्य", "ऑ", "ऑ", "औ", "ओ",
    "आ", "अ", "ई", "इ", "ऊ", "उ", "ऋ", "ऐ", "ए", "ॉ", "ू", "ु", "ं", "ा", "ृ", "्", "े",
    "ै", "ँ", "ी", "ः", "ो", "ौ"
]
array_two = [
    "…", "<", "क़", "ख़", "ग़", "ज़", "ड़", "ढ़", "फ़", "ç", "˜", ".", "'m", "]m",
    "Fmf", "Fm", ")", "!", "@", "#", "$", "%", "^", "&", "*", "(", "k|m", "em", "km",
    "Qm", "qm", "n", "¡", "¢", "1", "2", "4", ">", "?", "B", "I", "If", "Q", "ß", "q",
    "„", "‹", "•", "›", "§", "°", "¶", "¿", "Å", "Ë", "Ì", "Í", "Î", "Ý", "å", "6«", "7«",
    "8«", "9«", "|", "8Þ", "9Þ", "S", "s", "V", "v", "U", "u", "£", "3", "ª", "R", "r", "5",
    "H", "h", "‰", "´", "~", "`", "6", "7", "8", "9", "0", "0f", "T", "t", "Y", "y", "b",
    "W", "w", "G", "g", "K", "k", "ˆ", "A", "a", "E", "e", "D", "d", "o", "/", "N", "n", "J",
    "j", "Z", "z", "i", "if", ":", ";", "X", "x", "Ø", "cf‘", "c‘f", "cf}", "cf]", "cf", "c",
    "O{", "O", "pm", "p", "C", "P]", "P", "f‘", "\"", "'", "+", "f", "[", "\\", "]", "}", "F",
    "L", "M", "f]", "f}"
]

def convert_to_Preeti(modified_substring):
    processed_text = replace_symbols(modified_substring)
    print("Converted text: ", processed_text)
    return processed_text

def replace_symbols(modified_substring):
    if modified_substring != "":
        position_of_f = modified_substring.find("ि")
        while position_of_f != -1:
            character_left_to_f = modified_substring[position_of_f - 1]
            modified_substring = modified_substring.replace(
                character_left_to_f + "ि",
                "l" + character_left_to_f
            )
            position_of_f -= 1
            while (
                modified_substring[position_of_f - 1] == "्"
                and position_of_f != 0
            ):
                string_to_be_replaced = modified_substring[position_of_f - 2] + "्"
                modified_substring = modified_substring.replace(
                    string_to_be_replaced + "l",
                    "l" + string_to_be_replaced
                )
                position_of_f -= 2
            position_of_f = modified_substring.find("ि", position_of_f + 1)

        set_of_matras = "ािीुूृेैोौं:ँॅ"
        modified_substring += "  "
        space = " "
        position_of_half_R = modified_substring.find("र्")
        while position_of_half_R > 0:
            probable_position_of_Z = position_of_half_R + 2
            character_at_probable_position_of_Z = modified_substring[probable_position_of_Z]
            while character_at_probable_position_of_Z in set_of_matras:
                probable_position_of_Z += 1
                character_at_probable_position_of_Z = modified_substring[probable_position_of_Z]
            right_to_position_of_Z = probable_position_of_Z + 1
            if right_to_position_of_Z > 0:
                character_right_to_position_of_Z = modified_substring[right_to_position_of_Z]
                while "्" == character_right_to_position_of_Z:
                    probable_position_of_Z = right_to_position_of_Z + 1
                    character_at_probable_position_of_Z = modified_substring[probable_position_of_Z]
                    right_to_position_of_Z = probable_position_of_Z + 1
                    character_right_to_position_of_Z = modified_substring[right_to_position_of_Z]
            string_to_be_replaced = modified_substring[position_of_half_R + 2:probable_position_of_Z]
            modified_substring = modified_substring.replace(
                "र्" + string_to_be_replaced,
                string_to_be_replaced + "{"
            )
            position_of_half_R = modified_substring.find("र्")

        modified_substring = modified_substring[:-2]

        for input_symbol_idx in range(len(array_one)):
            idx = 0
            while idx != -1:
                modified_substring = modified_substring.replace(
                    array_one[input_symbol_idx],
                    array_two[input_symbol_idx]
                )
                idx = modified_substring.find(array_one[input_symbol_idx])

    return modified_substring

# customize as you wish to seem
def convert_file_to_preeti():
    # load document
    # read the content 
    # convert
    # save the document
    doc = Document('sample.docx') #input file name and path
    new_doc = Document()

    for para in doc.paragraphs:
        preeti_para = convert_to_Preeti(para.text)
        new_doc.add_paragraph(preeti_para)

    new_doc.save('okok.docx') #output file name and path

convert_to_Preeti('ओकोक्')
convert_file_to_preeti()